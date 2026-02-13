"""
Text Extraction Module
Extracts text content from various document formats (DOCX, PDF, PPTX, XLSX).
"""

import logging
import re
from pathlib import Path
from typing import Optional, Dict, Any

# Document processing libraries
import docx
import PyPDF2
import pdfplumber

from config import Config

# Set up logging
logger = logging.getLogger(__name__)


class TextExtractor:
    """Extracts text from various document formats."""
    
    @staticmethod
    def extract_from_docx(file_path: Path) -> Optional[str]:
        """
        Extract text from a DOCX file.
        
        Args:
            file_path: Path to DOCX file
        
        Returns:
            Optional[str]: Extracted text if successful, None otherwise
        """
        try:
            logger.info(f"Extracting text from DOCX: {file_path.name}")
            
            doc = docx.Document(file_path)
            
            # Extract text from paragraphs
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            
            # Extract text from tables
            table_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        table_text.append(" | ".join(row_text))
            
            # Combine all text
            all_text = paragraphs + table_text
            extracted_text = "\n\n".join(all_text)
            
            logger.info(f"✅ Extracted {len(extracted_text)} characters from {file_path.name}")
            return extracted_text
        
        except Exception as e:
            logger.error(f"Failed to extract text from {file_path.name}: {e}")
            return None
    
    @staticmethod
    def extract_from_pdf_pypdf2(file_path: Path) -> Optional[str]:
        """
        Extract text from a PDF file using PyPDF2.
        
        Args:
            file_path: Path to PDF file
        
        Returns:
            Optional[str]: Extracted text if successful, None otherwise
        """
        try:
            logger.info(f"Extracting text from PDF (PyPDF2): {file_path.name}")
            
            extracted_text = []
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                num_pages = len(pdf_reader.pages)
                
                logger.debug(f"PDF has {num_pages} pages")
                
                for page_num in range(num_pages):
                    page = pdf_reader.pages[page_num]
                    text = page.extract_text()
                    
                    if text.strip():
                        extracted_text.append(text)
            
            result = "\n\n".join(extracted_text)
            logger.info(f"✅ Extracted {len(result)} characters from {file_path.name}")
            return result
        
        except Exception as e:
            logger.error(f"Failed to extract text from {file_path.name} with PyPDF2: {e}")
            return None
    
    @staticmethod
    def extract_from_pdf_pdfplumber(file_path: Path) -> Optional[str]:
        """
        Extract text from a PDF file using pdfplumber (more robust than PyPDF2).
        
        Args:
            file_path: Path to PDF file
        
        Returns:
            Optional[str]: Extracted text if successful, None otherwise
        """
        try:
            logger.info(f"Extracting text from PDF (pdfplumber): {file_path.name}")
            
            extracted_text = []
            
            with pdfplumber.open(file_path) as pdf:
                num_pages = len(pdf.pages)
                logger.debug(f"PDF has {num_pages} pages")
                
                for page_num, page in enumerate(pdf.pages):
                    text = page.extract_text()
                    
                    if text and text.strip():
                        extracted_text.append(text)
                    else:
                        logger.debug(f"Page {page_num + 1} has no extractable text")
            
            result = "\n\n".join(extracted_text)
            logger.info(f"✅ Extracted {len(result)} characters from {file_path.name}")
            return result
        
        except Exception as e:
            logger.error(f"Failed to extract text from {file_path.name} with pdfplumber: {e}")
            return None
    
    @staticmethod
    def extract_from_pdf(file_path: Path, use_pdfplumber: bool = True) -> Optional[str]:
        """
        Extract text from a PDF file (tries pdfplumber first, falls back to PyPDF2).
        
        Args:
            file_path: Path to PDF file
            use_pdfplumber: Whether to try pdfplumber first
        
        Returns:
            Optional[str]: Extracted text if successful, None otherwise
        """
        if use_pdfplumber:
            # Try pdfplumber first (more robust)
            text = TextExtractor.extract_from_pdf_pdfplumber(file_path)
            
            # If pdfplumber fails or returns empty, try PyPDF2
            if not text or len(text.strip()) == 0:
                logger.info("pdfplumber returned no text, trying PyPDF2...")
                text = TextExtractor.extract_from_pdf_pypdf2(file_path)
        else:
            # Try PyPDF2 first
            text = TextExtractor.extract_from_pdf_pypdf2(file_path)
            
            # If PyPDF2 fails, try pdfplumber
            if not text or len(text.strip()) == 0:
                logger.info("PyPDF2 returned no text, trying pdfplumber...")
                text = TextExtractor.extract_from_pdf_pdfplumber(file_path)
        
        return text
    
    @staticmethod
    def normalize_text(text: str) -> str:
        """
        Normalize extracted text (remove extra whitespace, fix encoding issues).
        
        Args:
            text: Raw extracted text
        
        Returns:
            str: Normalized text
        """
        if not text:
            return ""
        
        # Replace multiple spaces with single space
        text = re.sub(r' +', ' ', text)
        
        # Replace multiple newlines with double newline
        text = re.sub(r'\n\n+', '\n\n', text)
        
        # Remove leading/trailing whitespace from each line
        lines = [line.strip() for line in text.split('\n')]
        text = '\n'.join(lines)
        
        # Remove any remaining extra whitespace
        text = text.strip()
        
        return text
    
    @staticmethod
    def extract_text(file_path: Path) -> Optional[str]:
        """
        Extract text from a file based on its extension.
        
        Args:
            file_path: Path to file
        
        Returns:
            Optional[str]: Extracted and normalized text if successful, None otherwise
        """
        if not file_path.exists():
            logger.error(f"File not found: {file_path}")
            return None
        
        # Get file extension
        file_ext = file_path.suffix.lower().lstrip('.')
        
        # Extract text based on file type
        text = None
        
        if file_ext == 'docx':
            text = TextExtractor.extract_from_docx(file_path)
        elif file_ext == 'pdf':
            text = TextExtractor.extract_from_pdf(file_path)
        else:
            logger.warning(f"Unsupported file type: {file_ext}")
            return None
        
        # Normalize text
        if text:
            text = TextExtractor.normalize_text(text)
        
        return text
    
    @staticmethod
    def save_extracted_text(
        text: str,
        output_path: Path,
        metadata: Optional[Dict[str, Any]] = None
    ) -> bool:
        """
        Save extracted text to a file with optional metadata.
        
        Args:
            text: Extracted text
            output_path: Path to save file
            metadata: Optional metadata to include at the top
        
        Returns:
            bool: True if successful, False otherwise
        """
        try:
            # Create output directory if it doesn't exist
            output_path.parent.mkdir(parents=True, exist_ok=True)
            
            with open(output_path, 'w', encoding='utf-8') as f:
                # Write metadata if provided
                if metadata:
                    f.write("="*60 + "\n")
                    f.write("METADATA\n")
                    f.write("="*60 + "\n")
                    for key, value in metadata.items():
                        f.write(f"{key}: {value}\n")
                    f.write("="*60 + "\n\n")
                
                # Write extracted text
                f.write(text)
            
            logger.info(f"✅ Saved extracted text to: {output_path}")
            return True
        
        except Exception as e:
            logger.error(f"Failed to save extracted text to {output_path}: {e}")
            return False
    
    @staticmethod
    def extract_and_save(
        file_path: Path,
        output_dir: Optional[Path] = None,
        metadata: Optional[Dict[str, Any]] = None
    ) -> Optional[Path]:
        """
        Extract text from a file and save it.
        
        Args:
            file_path: Path to source file
            output_dir: Directory to save extracted text (defaults to Config.EXTRACTED_TEXT_DIR)
            metadata: Optional metadata to include
        
        Returns:
            Optional[Path]: Path to saved text file if successful, None otherwise
        """
        if output_dir is None:
            output_dir = Config.EXTRACTED_TEXT_DIR
        
        # Extract text
        text = TextExtractor.extract_text(file_path)
        
        if not text:
            logger.error(f"No text extracted from {file_path.name}")
            return None
        
        # Create output filename (replace extension with .txt)
        output_filename = file_path.stem + ".txt"
        output_path = output_dir / output_filename
        
        # Add file info to metadata
        if metadata is None:
            metadata = {}
        
        metadata.update({
            "source_file": file_path.name,
            "file_type": file_path.suffix.lower(),
            "text_length": len(text),
            "word_count": len(text.split())
        })
        
        # Save extracted text
        if TextExtractor.save_extracted_text(text, output_path, metadata):
            return output_path
        else:
            return None


def test_text_extraction():
    """Test the text extraction module."""
    logging.basicConfig(
        level=logging.INFO,
        format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
    )
    
    print("\n🧪 Testing Text Extraction Module\n")
    
    # Check if there are any files to test with
    download_dir = Config.DOWNLOAD_DIR
    
    if not download_dir.exists():
        print(f"❌ Download directory not found: {download_dir}")
        print("   Please download some files first using sharepoint_api.py")
        return
    
    # Find test files
    docx_files = list(download_dir.glob("*.docx"))
    pdf_files = list(download_dir.glob("*.pdf"))
    
    test_files = docx_files[:2] + pdf_files[:2]  # Test up to 2 of each type
    
    if not test_files:
        print(f"❌ No DOCX or PDF files found in {download_dir}")
        print("   Please download some files first")
        return
    
    print(f"Found {len(test_files)} test files\n")
    
    # Test extraction
    Config.create_directories()
    
    for file_path in test_files:
        print(f"📄 Processing: {file_path.name}")
        
        output_path = TextExtractor.extract_and_save(file_path)
        
        if output_path:
            print(f"   ✅ Saved to: {output_path.name}\n")
        else:
            print(f"   ❌ Failed to extract text\n")
    
    print("✅ Text extraction test completed!\n")


if __name__ == "__main__":
    test_text_extraction()
